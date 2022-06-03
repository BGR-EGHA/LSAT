import docx
import logging
import traceback
import math
import matplotlib
from matplotlib.figure import Figure
import numpy as np
import os
from PyQt5.QtCore import QObject
import time
from core.libs.GDAL_Libs.Layers import Feature, Raster
from core.libs.Reporting.shared_report import shared_report
from core.widgets.ResultViewer.plotViewer_main import PlotFunctions

if os.name == "nt":
    prop = matplotlib.font_manager.FontProperties(fname="C:\\Windows\\Fonts\\Msyh.ttc")
else:  # ubuntu
    prop = matplotlib.font_manager.FontProperties(
        fname="/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc")


class woe_report(QObject):
    """
    Generates a Word document with information about the WoE calculation of a specific layer.
    """

    def __init__(self, project_path, result, name):
        super().__init__()
        try:
            # We open and prepare the template for editing
            doc = shared_report.opentemplate()
            # Load raster and feature into memory
            raster, feature = self.loadintomemory(result)
            # Basic info about the project and the calculationn
            doc = self.writebasics(doc, project_path, result)
            # info about the input data (feature and raster)
            doc = self.writeinputs(doc, result, feature, raster)
            # info about the analysis
            doc = self.writeanalysis(doc, result, raster, project_path)
            # Finally we save the document and tell the user where he can find it.
            shared_report.savedoc(name, project_path, doc, "WoE")
            logging.info(self.tr("Report saved under {}").format(
                os.path.join(project_path, "results", "WoE", "reports", name + ".docx")))
        except BaseException:
            tb = traceback.format_exc()
            logging.error(tb)

    def loadintomemory(self, result):
        """
        Loads the Raster and Feature into memory and returns both.
        If there are multiple features (subsamples) only the last one gets loaded.
        """
        raster = Raster(result["source"][0])
        feature = Feature(result["source"][1])
        return raster, feature

    def writebasics(self, doc, project_path, result):
        """
        Writes initial paragraph containing a table about the Project, Date and time
        """
        doc.add_heading(self.tr("Analysis Report"))
        doc.add_paragraph(
            self.tr("This is an automatically generated LSAT weight of evidence (WoE) analysis report."))
        doc.add_heading(self.tr("General Information"), level=2)
        self._add_basictable(doc, project_path, result)
        return doc

    def _add_basictable(self, doc, project_path, result):
        """
        Gets called by writebasics. Adds a basic info table
        """
        basictable = doc.add_table(5, 2, style='Light Shading Accent 1')
        basictable.cell(0, 0).text = self.tr("Project name")
        # Name of the lowest folder is the project name
        # this should be \ (windows os.sep) but isn't...
        basictable.cell(0, 1).text = project_path.split("/")[-1]
        basictable.cell(1, 0).text = self.tr("Time and date")
        basictable.cell(1, 1).text = time.strftime("%Y-%m-%d %H:%M", time.localtime())
        basictable.cell(2, 0).text = self.tr("Subsampling type")
        basictable.cell(2, 1).text = self._get_subsample_type(result["metadata"][3])
        basictable.cell(3, 0).text = self.tr("Sample size [%]")
        basictable.cell(3, 1).text = result["metadata"][2]  # Sample size
        basictable.cell(4, 0).text = self.tr("Number of subsamples")
        basictable.cell(4, 1).text = result["metadata"][4]  # Subsample count
        try:
            result["metadata"][5] # Random Seed
            basictable.add_row()
            basictable.cell(5, 0).text = self.tr("Seed used to Initialize random")
            basictable.cell(5, 1).text = result["metadata"][5]  # Random Seed
        except IndexError: # only on the fly subsampling features a seed
            pass
        return doc

    def _get_subsample_type(self, sampleNrfromMetadata: str) -> str:
        """
        Gets called by _add_basictable
        The type of subsampling used gets stored in the npz file. We generate a string for the
        report based on that number.
        """
        if sampleNrfromMetadata == "1":
            Subsampling = "1 - No subsampling"
        elif sampleNrfromMetadata == "2":
            Subsampling = "2 - On-the-fly subsampling"
        elif sampleNrfromMetadata == "3":
            Subsampling = "3 - Predefined Subsamples"
        return Subsampling

    def writeinputs(self, doc, result, feature, raster):
        """
        Writes paragraph with table for the input feature and input raster
        """
        doc.add_heading(self.tr("Input Data"), level=2)
        doc.add_heading(self.tr("Input feature"), level=3)
        self._add_featuretable(doc, result, feature)
        doc.add_heading(self.tr("Input raster"), level=3)
        self._add_rastertable(doc, result, raster)
        if raster.rat is not None:
            doc.add_heading(self.tr("Raster Attribute Table"), level=4)
            self._add_rat(doc, result, raster)
        return doc

    def _add_featuretable(self, doc, result, feature):
        """
        Gets called by writeinputs.
        """
        featuretable = doc.add_table(3, 2, style='Light Shading Accent 1')
        featuretable.cell(0, 0).text = self.tr("Source")
        featuretable.cell(0, 1).text = self._get_featurepath(result)
        featuretable.cell(1, 0).text = self.tr("Geometry")
        featuretable.cell(1, 1).text = feature.geometryName
        featuretable.cell(2, 0).text = self.tr("Number of Events")
        featuretable.cell(2, 1).text = self._get_eventcount(result, feature)

    def _get_featurepath(self, result) -> str:
        """
        Gets called by _add_featuretable.
        Returns the path to the input feature if there are no (1) subsamples, or the path to the
        folder where multiple samples are stored in. Either is a string.
        """
        if int(result["metadata"][4]) == 1:
            path = result["source"][1]
        else:
            if int(result["metadata"][3]) == 2:  # On-the-fly Subsamples
                path = self.tr("Multiple on-the-fly subsampled from {}").format(result["source"][1])
            elif int(result["metadata"][3]) == 3:  # Predefined Subsamples
                path = self.tr("Multiple in {}").format(os.path.split(result["source"][1])[0])
        return path

    def _get_eventcount(self, result, feature) -> str:
        """
        Gets called by _add_featuretable.
        Returns a string with number of events in the input feature. If there multiple subsamples
        we let the user know that it is only about the last analyzed sample
        """
        if int(result["metadata"][4]) == 1:
            count = str(feature.count)
        else:
            count = self.tr("Last analyzed: {}").format(feature.count)
        return count

    def _add_rastertable(self, doc, result, raster):
        """
        Gets called by writeinputs.
        """
        rastertable = doc.add_table(9, 2, style='Light Shading Accent 1')
        rastertable.cell(0, 0).text = self.tr("Source")
        rastertable.cell(0, 1).text = result["source"][0]
        rastertable.cell(1, 0).text = self.tr("Dimensions (cols, rows)")
        rastertable.cell(1, 1).text = "({} , {})".format(raster.cols, raster.rows)
        rastertable.cell(2, 0).text = self.tr("Minimum")
        rastertable.cell(2, 1).text = str(round(raster.min, 2))
        rastertable.cell(3, 0).text = self.tr("Maximum")
        rastertable.cell(3, 1).text = str(round(raster.max, 2))
        rastertable.cell(4, 0).text = self.tr("Mean")
        rastertable.cell(4, 1).text = self.tr(str(round(raster.stats[2], 3)))
        rastertable.cell(5, 0).text = self.tr("Std")
        rastertable.cell(5, 1).text = self.tr(str(round(raster.stats[3], 3)))
        rastertable.cell(6, 0).text = self.tr("Cell size [x, y]")
        rastertable.cell(6, 1).text = str(raster.cellsize)
        rastertable.cell(7, 0).text = self.tr("NoData")
        rastertable.cell(7, 1).text = str(raster.nodata)
        rastertable.cell(8, 0).text = self.tr("Number of Unique Values")
        rastertable.cell(8, 1).text = str(raster.uniqueValuesCount)

    def _add_rat(self, doc, result, raster):
        """
        Adds the Raster Attribute Table.
        Returns Name of Columns in RAT.
        """
        values = raster.rat.ReadAsArray(0).tolist()
        rat = raster.getRasterAttributeTableFromBand()  # assuming band 1
        ratcolnames = []
        for i in range(rat.GetColumnCount()):
            ratcolnames.append(rat.GetNameOfCol(i))
        rattable = doc.add_table(
            len(values) + 1,
            len(ratcolnames),
            style='Medium Shading 1 Accent 1')
        for name in range(len(ratcolnames)):
            values = rat.ReadAsArray(name).tolist()
            for row, value in enumerate(values):
                if row == 0:
                    cell = rattable.cell(row, name)
                    cell.text = ratcolnames[name]
                cell = rattable.cell(row + 1, name)
                if isinstance(value, bytes):  # RAT uses Byte strings
                    try:
                        value = value.decode("utf-8")
                    except UnicodeDecodeError:
                        value = value.decode("latin-1")
                cell.text = str(value)

    def writeanalysis(self, doc, result, raster, project_path):
        """
        Writes a paragraph with a plot showing the results and another Raster Attribute Table this
        time we add the calculated weight and sort it with that.
        """
        doc.add_heading(self.tr("Analysis Results"), level=2)
        doc.add_heading(self.tr("Visualized Results"), level=3)
        self._add_visresultsplot(doc, result, project_path)
        doc.add_paragraph(self.tr(
            "Fig. 1 Results overview: a) Class distribution; b) Landslide in class distribution; c) Weight distribution; d) Analysis ROC curve"))
        doc.add_heading(self.tr("Raster values sorted by weight"), level=3)
        doc.add_paragraph(self.tr(
            "The significance has two directions. In general weights > 1 are significantly positive and weights < -1 are significantly negative."))
        self._add_tableweight(doc, result, raster)
        self._add_resultstablesection(doc, result)
        return doc

    def _add_visresultsplot(self, doc, result, project_path):
        """
        Adds Matplotlib graphics as seen in the Result Viewer.
        """
        fig = Figure()

        labels = PlotFunctions.getLabels(result)
        self._add_plothistogram(fig, result, labels)
        self._add_plotlandslidedistribution(fig, result, labels)
        self._add_plotweightdistribution(fig, result, labels)
        self._add_plotroc(fig, result)
        fig.set_size_inches(8, 6)
        fig.tight_layout()

        tmp_path = os.path.join(project_path, "workspace", "tmp_fig.png")
        fig.savefig(tmp_path, dpi = 300)
        doc.add_picture(tmp_path, width=docx.shared.Cm(15))
        os.remove(tmp_path)

    def _add_plothistogram(self, fig, result, labels):
        """
        Adds a subplot to fig with a histogram of raster values.
        """
        histogram = fig.add_subplot(221)
        histogram.text(0.03, 0.97, "a)", transform=histogram.transAxes,
                    fontsize=12, va='top')
        histogram.set_xlabel(self.tr('Classes'), fontproperties=prop)
        histogram.set_ylabel(self.tr('Class pixel count'), fontproperties=prop)
        histogram.set_ylim(0, result["tab"]['Class'].max() * 1.1)
        histogram.set_xlim(0, len(result["tab"]) + 1)
        y = []
        x = []
        for i in range(len(result["tab"])):
            y.append(result["tab"]['Class'][i])
            x.append(i + 1)
        if len(x) < 10:
            histogram.set_xticks(x)
            histogram.set_xticklabels(labels)
        histogram.bar(x, y, align='center')

    def _add_plotlandslidedistribution(self, fig, result, labels):
        """
        Adds a subplot to fig showing the landslide distribution in class
        """
        lsdistribution = fig.add_subplot(222)
        lsdistribution.text(0.03, 0.97, "b)", transform=lsdistribution.transAxes,
                       fontsize=12, va='top')
        lsdistribution.set_xlabel(self.tr('Classes'), fontproperties=prop)
        lsdistribution.set_ylabel(self.tr('Landslide pixel count'), fontproperties=prop)
        lsdistribution.set_xlim(0, len(result["tab"]) + 1)
        if len([key for key in result]) <= 6:  # multiple samples check
            lsdistribution.set_ylim(0, result["tab"]['Landslides'].max() * 1.1)
            y = []
            x = []
            for i in range(len(result["tab"])):
                y.append(result["tab"]['Landslides'][i])
                x.append(i + 1)
            if len(x) < 10:
                lsdistribution.set_xticks(x)
                lsdistribution.set_xticklabels(labels)
            lsdistribution.bar(x, y, align='center')
        else:
            if len(result["tab"]) < 10:
                lsdistribution.set_xticks(range(1, len(result["tab"]) + 1))
                lsdistribution.set_xticklabels(labels)
            data = []
            for i in range(len(result["tab"])):
                data.append(result["landslides"][i])
            violin = lsdistribution.violinplot(data, showmedians=True)

    def _add_plotweightdistribution(self, fig, result, labels):
        """
        Adds a plot to fig showing the distribution of the weights
        """
        weightdistribution = fig.add_subplot(223)
        weightdistribution.text(0.03, 0.97, "c)", transform=weightdistribution.transAxes,
                            fontsize=12, va='top')
        weightdistribution.set_xlabel(self.tr('Classes'), fontproperties=prop)
        weightdistribution.set_ylabel(self.tr('Total weight'), fontproperties=prop)
        weightdistribution.set_xlim(0, len(result["tab"]) + 1)
        if len([key for key in result]) <= 6:  # multiple samples check
            y = []
            x = []
            Std = []
            colour = []
            for i in range(len(result["tab"])):
                y.append(result["tab"]['Weight'][i])
                x.append(i + 1)
                Std.append(math.sqrt(result["tab"]['Variance'][i]))
                if result["tab"]['Weight'][i] <= 0.0 and result["tab"]['Weight'][i] >= -1.0:
                    a = 'y'
                elif result["tab"]['Weight'][i] > 0.0 and result["tab"]['Weight'][i] <= 1.0:
                    a = 'y'
                elif result["tab"]['Weight'][i] < -1.0:
                    a = 'r'
                else:
                    a = 'g'
                colour.append(a)
            if len(x) < 10:
                weightdistribution.set_xticks(x)
                weightdistribution.set_xticklabels(labels)
            weightdistribution.bar(x, y, color=colour, yerr=Std, align='center')
        else:
            if len(result["tab"]) < 10:
                weightdistribution.set_xticks(range(1, len(result["tab"]) + 1))
                weightdistribution.set_xticklabels(labels)
            data = []
            colors = []
            for i in range(len(result["tab"])):
                data.append(result['weight'][i])
                if result["tab"]['Weight'][i] <= 0.0 and result["tab"]['Weight'][i] >= -1.0:
                    a = 'y'
                elif result["tab"]['Weight'][i] > 0.0 and result["tab"]['Weight'][i] <= 1.0:
                    a = 'y'
                elif result["tab"]['Weight'][i] < -1.0:
                    a = 'r'
                else:
                    a = 'g'
                colors.append(a)
            violin = weightdistribution.violinplot(data, showmeans=False, showmedians=True)
            for patch, color in zip(violin['bodies'], colors):
                patch.set_facecolor(color)

    def _add_plotroc(self, fig, result):
        """
        Adds a ROC Curve to fig.
        """
        rocplot = fig.add_subplot(224)
        rocplot.text(0.03, 0.97, "d)", transform=rocplot.transAxes,
                                fontsize=12, va='top')
        rocplot.set_xlabel(self.tr('False Positive Rate (1-Specificity)'),
                           fontproperties=prop)
        rocplot.set_ylabel(self.tr('True Positive Rate (Sensitivity)'),
                           fontproperties=prop)
        x0 = [0.0, 1.0]
        y0 = [0.0, 1.0]
        rocplot.set_ylim(0, 1.0)
        rocplot.set_xlim(0, 1.0)
        if len([key for key in result]) <= 6:  # check if multiple samples
            x = result['roc_x']
            y = result['roc_y']
            auc = result['auc']
            rocplot.plot(x, y, label=self.tr("AUC: %.2f") % (float(auc)), color="red", linewidth=2)
        else:
            sortTable = np.sort(result["tab"], order=['Weight'])[::-1]
            x = [0.0]
            y = [0.0]
            for i in range(len(result["tab"])):
                x.append(x[i] + float(sortTable['Class'][i]) / float(sortTable['Class'].sum()))
                y.append(y[i] + float(sortTable['Landslides'][i]) /
                         float(sortTable['Landslides'].sum()))
            unique = np.unique(result["roc_x"].ravel())
            y_all = []
            for i in range(len(result["roc_x"])):
                new_y = np.interp(unique, result["roc_x"][i], result["roc_y"][i])
                y_all.append(new_y)
            max_y = []
            min_y = []
            mean_y = []
            median_y = []
            for i in np.transpose(y_all):
                min_y.append(min(i))
                max_y.append(max(i))
                mean_y.append(np.mean(i))
                median_y.append(np.median(i))
            auc = []
            for i in range(len(unique) - 1):
                auc.append((unique[i + 1] - unique[i]) * mean_y[i + 1] - (
                    (unique[i + 1] - unique[i]) * (mean_y[i + 1] - mean_y[i]) / 2))

            rocplot.fill_between(unique, min_y, max_y, color="grey", linestyle="--", alpha=0.5,
                                 label=self.tr("range: %.2f - %.2f") % (
                                     result['auc'].max(), result['auc'].min()))
            rocplot.plot(unique, mean_y, color="red", linewidth=2,
                         label=self.tr("mean: %.2f") % (sum(auc)))
            rocplot.plot(x0, y0, label=self.tr("prior: 0.5"), color="blue")
            rocplot.legend(loc="lower right")

    def _add_tableweight(self, doc, result, raster):
        """
        Adds a table with raster values sorted by calculated Weight.
        """
        rat = raster.getRasterAttributeTableFromBand()  # assuming band 1
        ratcolnames = []
        for i in range(rat.GetColumnCount()):
            ratcolnames.append(rat.GetNameOfCol(i))
        try:  # Case sensitive
            value_index = ratcolnames.index("VALUE")
        except ValueError:
            value_index = ratcolnames.index("Value")
        values = rat.ReadAsArray(value_index).tolist()
        if raster.nodata in values:
            values.remove(raster.nodata)
        weighttable = doc.add_table(len(values) + 1, 3, style='Medium Shading 1 Accent 1')
        rankarray = np.zeros(shape=(len(values),), dtype=[(self.tr("Raster value"), "i"),
                                                          (self.tr("Weight"), "f"),
                                                          (self.tr("STD"), "f")])
        for i in range(len(values)):
            rankarray[self.tr("Raster value")][i] = values[i]
            rankarray[self.tr("Weight")][i] = round(result["tab"]["Weight"][i], 3)
            rankarray[self.tr("STD")][i] = round(np.sqrt(result["tab"]["Variance"][i]), 3)
        rankarray[::-1].sort(order=self.tr("Weight"))
        for i in range(len(values) + 1):
            for j in range(len(rankarray.dtype.names)):
                cell = weighttable.cell(i, j)
                if i == 0:
                    cell.text = rankarray.dtype.names[j]
                else:
                    cell.text = str(rankarray[i - 1][j])

    def _add_resultstablesection(self, doc, result):
        """
        Adds a table with all the information we gathered for Weights of Evidence. Because there
        are a lot of columns we use the paper horizontally.
        """
        resulttablesection = doc.add_section()
        landscape_width = resulttablesection.page_height
        landscape_height = resulttablesection.page_width
        resulttablesection.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        resulttablesection.page_width = landscape_width
        resulttablesection.page_height = landscape_height
        doc.add_heading(self.tr("All Results"), level=3)
        resulttable = doc.add_table(len(result["tab"]) + 1, len(result["tab"].dtype.names),
                                    style='Medium Shading 1 Accent 1')
        for i, name in enumerate(result["tab"].dtype.names):
            for j in range(len(result["tab"]) + 1):
                cell = resulttable.cell(j, i)
                if j == 0:
                    cell.text = result["tab"].dtype.names[i]
                else:
                    cell.text = str(result["tab"][name][j - 1])
        # https://stackoverflow.com/questions/43007725/python-docx-how-to-change-table-font-size
        for row in resulttable.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        run.font.size = docx.shared.Pt(7.5)
